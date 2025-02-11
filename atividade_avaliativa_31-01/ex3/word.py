import docx
import lorem
# Exercício 1

doc = docx.Document()
runner = doc.add_heading().add_run('Salve')

runner.bold = True
runner.underline = True

p = lorem.paragraph()

doc.add_paragraph(p)
doc.add_paragraph(p)
doc.save('atividade_avaliativa_31-01/ex3/documento.docx')


# Exercício 2


doc = docx.Document('atividade_avaliativa_31-01/ex3/documento.docx')
for paragraph in doc.paragraphs:
    print(paragraph.text)


# Exercício 3


doc = docx.Document('atividade_avaliativa_31-01/ex3/documento.docx')
palavras = 0
for paragraph in doc.paragraphs:
    palavras += len(paragraph.text.split())

print(palavras)


# Exercício 4


doc = docx.Document()

table = doc.add_table(rows=5, cols=3)

for coluna in range(len(table.columns)):
    for linha in range(5):
        table.cell(linha, coluna).text = (f'Linha {linha + 1}, Coluna {coluna + 1}')
       

doc.save('atividade_avaliativa_31-01/ex3/tabela.docx')


# Exercício 5

doc = docx.Document('atividade_avaliativa_31-01/ex3/Python.docx') 
for paragraph in doc.paragraphs:
    if 'Python' in paragraph.text:
        paragraph.text = paragraph.text.replace('Python', 'Elixir')

doc.save('atividade_avaliativa_31-01/ex3/Elixir.docx')