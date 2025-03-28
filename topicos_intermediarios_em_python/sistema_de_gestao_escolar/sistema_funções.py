import re
import csv
from docx import Document
from openpyxl import Workbook

def menu():
    print(f'\n======SISTEMA DE GESTÃO ESCOLAR EM TERMINAL======\n')
    print(f'1 - Adicionar Aluno')
    print(f'2 - Remover Aluno')
    print(f'3 - Listar Alunos')
    print(f'4 - Gerar Relatório')
    print(f'5 - Gerar Planilha')
    print(f'0 - Sair\n')
    print(f'==================================================')


def função_escolha(min, max):
    while True:
        try:
            escolha = int(input(f'\nDigite o número da opção que deseja realizar: '))
            while escolha < min or escolha > max:
                print(f'\nNúmero inválido! Digite um número entre {min} e {max}.')
                escolha = int(input(f'\nDigite o número da opção que deseja realizar: '))
            return escolha
        except ValueError:
            print(f'\nEntrada inválida! Tente novamente.')

def adicionar_aluno(lista):
    print(f'\nPreencha as informações do aluno para adicionar-lo ao sistema.')

    nome = input(f'\nNome: ')
    while True:
        if validar_nome(nome):
            print(f'\nNome válido')
            break
        nome = input(f'\nNome inválido, tente novamente: ')
            
    matricula = input(f'\nMatricula, (XXXXXXXXX): ')
    while True:
        if validar_matricula(matricula, lista):
            print(f'\nMatricula válida!')
            break
        matricula = input(f'\nTente novamente: ')
        
    data = input(f'\nData de Nascimento, (XX/XX/XXXX): ')
    while True:
        if validar_data(data):
            print(f'\nData válida!')
            break
        else:
            data = input(f'\nData inválida, tente novamente: ')

    msg = f'\nQuantas disciplinas devemos adicionar no registro desse aluno, (Mínimo de 3 disciplinas): '
    qnt_disc = função_try_int( msg)
    while qnt_disc < 3:
        print(f'\nNo mínimo 3 disciplinas! ')
        qnt_disc = função_try_int(msg)

    disciplinas = ['Português', 'Matemática', 'Geografia', 'História', 'Física', 'Química', 'Biologia']

    notas = {}

    for i in range(qnt_disc):
        for n, disc in enumerate(disciplinas):
            print(f'{n + 1} - {disc}')

        msg = f'\nOpção da {i + 1}º disciplina: '
        disc = verifica_disc(disciplinas, msg)
        
        msg = f'\nNota de {disciplinas[disc - 1]}: '
        nota = função_try_float(msg)
        while nota < 0 or nota > 10:
            print(f'\nA nota deve ser entre 0 e 10! ')
            nota = função_try_float(msg)
        
        notas[disciplinas[disc - 1]] = nota
    
        del disciplinas[disc - 1]

    soma = 0
    for nota in notas.values():
        soma += nota
    media = round(soma/qnt_disc, 2)


    msg = f'\nQual a porcentagem de faltas desse aluno? '
    qnt_faltas = função_try_int(msg)
    while qnt_faltas < 0 or qnt_faltas > 100:
        print(f'\nDigite um valor válido! ')
        qnt_faltas = função_try_int(msg)


    lista.append([nome, matricula, data, notas, media, qnt_faltas])

    atualiza_arquivo(lista)


def remover_aluno(lista):
    while True:
        matricula = input(f'\nDigite a matricula do aluno, que deseja remover do sistema: ')
        achou = False
        for i in range(len(lista)):
            if matricula == lista[i][1]:
                del lista[i]
                print(f'\nAluno removido com sucesso!')
                atualiza_arquivo(lista)
                achou = True
                break
        if achou:
            break
        print(f'\nMatricula não encontrada! Tente novamente. ')

def listar_alunos(lista):
    for aluno in lista:
        print(f'{aluno[0]} - {aluno[1]} - {aluno[3]} - {aluno[5]}% de Faltas.')
        


def validar_data(data):
    pattern = "(0[1-9]|[12][0-9]|3[01])/(0[1-9]|1[012])/(19[5-9][0-9]|202[0-5]|20[0-1][0-9])"

    if re.search(pattern, data):
        return True
    else:
        return False
    
def validar_matricula(matricula, lista):
    if any(chr.isalpha() for chr in matricula) or len(matricula) != 9:
        print(f'\nMatricula inválida!')
        return False
    for i in range(len(lista)):
        if matricula == lista[i][1]:
            print(f'\nMatricula já esta em uso!')
            return False
    return True
    
def validar_nome(nome):
    if any(chr.isdigit() for chr in nome):
        return False
    return True

def função_try_int(msg):
    while True:
        try:
            var = int(input(msg))
            return var
        except ValueError:
            print(f'\nEntrada inválida, tente novamente. ')


def função_try_float(msg):
    while True:
        try:
            var = float(input(msg))
            return var
        except ValueError:
            print(f'\nEntrada inválida, tente novamente. ')


def verifica_disc(lista, msg):
    disc = função_try_int(msg)
    while (disc < 1 or disc > len(lista)):
        print(f'\nEscolha uma opção entre 1 e {len(lista)}')
        disc = função_try_int(msg)

    return disc

def matricula_unica(matricula, lista):
    for i in len(lista):
        if matricula == lista[i][1]:
            return False
        
    return True


def atualiza_arquivo(lista):
    with open('sistema_de_gestao_escolar/alunos.csv', 'w', newline="") as arquivo:
        escritor = csv.writer(arquivo)
        for i in range(len(lista)):
            escritor.writerow(lista[i])


def relatorio_word_unico(matricula, lista):
    for i in range(len(lista)):
        if matricula == lista[i][1]:
            break
    
    doc = Document()

    doc.add_heading(f'Relatório de {lista[i][0]}', level=0)
    doc.add_paragraph(f'Nome: {lista[i][0]}')
    doc.add_paragraph(f'Matricula: {lista[i][1]}')
    doc.add_paragraph(f'Data de Nascimento: {lista[i][2]}')
    doc.add_paragraph(f'Notas: {lista[i][3]}')
    doc.add_paragraph(f'Média: {lista[i][4]}')
    doc.add_paragraph(f'Faltas: {lista[i][5]}%')
    if lista[i][4] < 7 or lista[i][5] > 25:
        doc.add_paragraph(f'Status de aprovação: Reprovado')
    else:
        doc.add_paragraph(f'Status de aprovação: Aprovado')

    doc.save(f'sistema_de_gestao_escolar/{lista[i][1]}_relatorio.docx')
    print(f'Relatório criado com sucesso! ')
    
def relatorio_word_grupo(lista):
    doc = Document()

    doc.add_heading('Relatório dos Alunos', level=0)

    for aluno in lista:
        doc.add_heading(f'Relatório de {aluno[0]}', level=1)
        doc.add_paragraph(f'Nome: {aluno[0]}')
        doc.add_paragraph(f'Matricula: {aluno[1]}')
        doc.add_paragraph(f'Data de Nascimento: {aluno[2]}')
        doc.add_paragraph(f'Notas: {aluno[3]}')
        doc.add_paragraph(f'Média: {aluno[4]}')
        doc.add_paragraph(f'Faltas: {aluno[5]}%')

        print(f'{aluno[4]}')
        if float(aluno[4]) < 7 or float(aluno[5]) > 25:
            doc.add_paragraph(f'Status de aprovação: Reprovado')
        else:
            doc.add_paragraph(f'Status de aprovação: Aprovado')
    
    doc.save(f'sistema_de_gestao_escolar/relatorio_grupo.docx')
    print(f'\nRelatório do grupo criado com sucesso! ')

def planilha_excel(lista):
    wb = Workbook()
    ws = wb.active
    ws.title = 'Alunos'

    cabechalos = ['Nome', 'Matricula', 'Data de Nascimento', 'Notas', 'Média', 'Porcentagem de Faltas', 'Status de Aprovação']
    ws.append(cabechalos)

    for aluno in lista:
        status = "Aprovado" if (float(aluno[4]) >= 7 and float(aluno[5]) <= 25) else "Reprovado"
        ws.append([aluno[0], aluno[1], aluno[2], aluno[3], aluno[4], aluno[5], status])
    
    wb.save('topicos_intermediarios_em_python_/sistema_de_gestao_escolar/alunos.xlsx')
    print(f'Planilha criada com sucesso! ')
